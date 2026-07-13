from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from wenmode_docx import DOCXRenderer, markdown_to_docx, save_markdown_as_docx

from wenmode import Wenmode
from wenmode.presets import github


def test_markdown_to_docx_writes_common_blocks() -> None:
    document = markdown_to_docx(
        '# Title\n\n'
        'Hello **strong** and *emphasis* with `code`.\n\n'
        '- first\n'
        '- [x] done\n\n'
        '> quoted\n\n'
        '| A | B |\n'
        '| --- | --- |\n'
        '| 1 | 2 |\n'
    )

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    tables = document.tables

    assert paragraphs[:5] == ['Title', 'Hello strong and emphasis with code.', 'first', '[x] done', 'quoted']
    assert tables[0].cell(0, 0).text == 'A'
    assert tables[0].cell(1, 1).text == '2'


def test_save_markdown_as_docx_creates_readable_file(tmp_path: Path) -> None:
    output = tmp_path / 'sample.docx'

    save_markdown_as_docx('## Sample\n\nBody text.\n', output)

    saved = Document(output)
    assert [paragraph.text for paragraph in saved.paragraphs] == ['Sample', 'Body text.']


def test_docx_renderer_can_be_used_with_wenmode() -> None:
    renderer = DOCXRenderer()

    document = Wenmode(github, renderer=renderer).render('# Rendered\n\nFrom a renderer.\n')

    assert [paragraph.text for paragraph in document.paragraphs] == ['Rendered', 'From a renderer.']


def test_docx_renderer_can_be_reused_without_sharing_document_state() -> None:
    renderer = DOCXRenderer()
    wen = Wenmode(github, renderer=renderer)

    first = wen.render('# First\n')
    second = wen.render('# Second\n')

    assert [paragraph.text for paragraph in first.paragraphs] == ['First']
    assert [paragraph.text for paragraph in second.paragraphs] == ['Second']
    assert first is not second


def test_docx_renderer_outputs_word_hyperlink_and_bottom_border() -> None:
    document = markdown_to_docx('Visit [Wenmode](https://wenmode.lepture.com).\n\n---\n')

    hyperlink = document.paragraphs[0]._p.find(qn('w:hyperlink'))
    border = document.paragraphs[1]._p.pPr.find(qn('w:pBdr')).find(qn('w:bottom'))

    assert hyperlink is not None
    assert hyperlink.find(qn('w:r')).text == 'Wenmode'
    assert document.paragraphs[0].text == 'Visit Wenmode.'
    assert border.get(qn('w:val')) == 'single'


def test_docx_renderer_uses_default_table_style_unless_configured() -> None:
    plain = markdown_to_docx('| A |\n| --- |\n| B |\n')
    styled = DOCXRenderer(table_style='Table Grid').render(Wenmode(github).parse('| A |\n| --- |\n| B |\n'))

    assert plain.tables[0].style.name == 'Normal Table'
    assert styled.tables[0].style.name == 'Table Grid'
