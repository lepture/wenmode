from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import cast
from urllib.parse import urlparse

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph as DocxParagraph

from wenmode import Wenmode
from wenmode.nodes import (
    Blockquote,
    Break,
    Code,
    Delete,
    Emphasis,
    FootnoteDefinition,
    FootnoteReference,
    Heading,
    Html,
    Image,
    InlineCode,
    Link,
    List,
    ListItem,
    Node,
    Paragraph,
    Parent,
    Root,
    Strong,
    Table,
    TableCell,
    TableRow,
    Text,
    ThematicBreak,
)
from wenmode.presets import github
from wenmode.renderers import BaseRenderer, RenderContext

LIST_INDENT = 0.5
MAX_INDENT = 5.5


def markdown_to_docx(source: str, document: DocumentObject | None = None) -> DocumentObject:
    """Convert Markdown source to a DOCX document object.

    The returned document is not saved automatically. Pass an existing
    `python-docx` document to append Markdown content to a template.
    """
    renderer = DOCXRenderer()
    context = DOCXRenderContext(document=document if document is not None else Document())
    root = Wenmode(github, renderer=renderer).parse(source)
    return renderer.render(root, context)


def save_markdown_as_docx(source: str, output: str | Path, document: DocumentObject | None = None) -> DocumentObject:
    """Convert Markdown source and save it to a DOCX file."""
    target = markdown_to_docx(source, document=document)
    target.save(output)
    return target


@dataclass
class DOCXRenderContext(RenderContext):
    """Render context for DOCX output."""

    document: DocumentObject = field(default_factory=Document)
    list_level: int = 0


class DOCXRenderer(BaseRenderer):
    """Render Wenmode nodes into a `python-docx` document."""

    name = 'docx'

    def __init__(self, table_style: str | None = None, paragraph_style: str | None = None) -> None:
        super().__init__()
        self.table_style = table_style
        self.paragraph_style = paragraph_style

    def create_context(self, node: Node | None = None) -> DOCXRenderContext:
        return DOCXRenderContext()

    def render(self, node: Node, context: RenderContext | None = None) -> DocumentObject:
        """Render a node and return the document created for this render call."""
        docx_context = cast(DOCXRenderContext, context) if context is not None else self.create_context(node)
        if node.type == 'root':
            self.render_root(node, docx_context)
        else:
            self.render_node(node, docx_context)
        return docx_context.document

    def render_unknown(self, node: Node, context: RenderContext) -> str:
        docx_context = cast(DOCXRenderContext, context)
        children = getattr(node, 'children', None)
        if isinstance(children, list):
            self.render_children(cast(list[Node], children), docx_context)
        else:
            value = getattr(node, 'value', None)
            if isinstance(value, str):
                docx_context.document.add_paragraph(value)
        return ''

    def render_list_item(self, node: ListItem, style: str, context: DOCXRenderContext) -> None:
        remaining = list(node.children)
        first = remaining.pop(0) if remaining else None
        paragraph = context.document.add_paragraph(style=style)
        paragraph.paragraph_format.left_indent = Inches(min((context.list_level + 1) * LIST_INDENT, MAX_INDENT))
        paragraph.paragraph_format.line_spacing = 1
        if node.checked is not None:
            paragraph.add_run('[x] ' if node.checked else '[ ] ')
        if isinstance(first, Paragraph):
            self.render_inlines(paragraph, first.children)
        elif first is not None:
            self.render_at_list_level(first, context, context.list_level + 1)

        for child in remaining:
            if isinstance(child, List):
                self.render_at_list_level(child, context, context.list_level + 1)
            else:
                self.render_node(child, context)

    def render_at_list_level(self, node: Node, context: DOCXRenderContext, list_level: int) -> None:
        previous_level = context.list_level
        context.list_level = list_level
        try:
            self.render_node(node, context)
        finally:
            context.list_level = previous_level

    def render_cell_content(self, node: TableCell, cell: _Cell, paragraph: DocxParagraph) -> None:
        block_children = [child for child in node.children if getattr(child, 'block', False)]
        if not block_children:
            self.render_inlines(paragraph, node.children)
            return

        for index, child in enumerate(block_children):
            target = paragraph if index == 0 else cell.add_paragraph()
            if isinstance(child, Paragraph):
                self.render_inlines(target, child.children)
            elif isinstance(child, Parent):
                self.render_inlines(target, child.children)

    def add_paragraph(
        self, context: DOCXRenderContext, children: list[Node] | None = None, style: str | None = None
    ) -> DocxParagraph:
        try:
            paragraph = context.document.add_paragraph(style=style or self.paragraph_style)
        except KeyError as error:
            style_name = style or self.paragraph_style
            raise ValueError(f'Unable to apply style {style_name}.') from error
        if children is not None:
            self.render_inlines(paragraph, children)
        return paragraph

    def render_inlines(self, paragraph: DocxParagraph, children: list[Node]) -> None:
        for child in children:
            self.render_inline(child, paragraph)

    def render_inline(
        self, node: Node, paragraph: DocxParagraph, bold: bool = False, italic: bool = False, strike: bool = False
    ) -> None:
        if isinstance(node, Text):
            run = paragraph.add_run(node.value)
            run.bold = bold
            run.italic = italic
            run.font.strike = strike
        elif isinstance(node, Strong):
            self.render_styled_children(node, paragraph, bold=True, italic=italic, strike=strike)
        elif isinstance(node, Emphasis):
            self.render_styled_children(node, paragraph, bold=bold, italic=True, strike=strike)
        elif isinstance(node, Delete):
            self.render_styled_children(node, paragraph, bold=bold, italic=italic, strike=True)
        elif isinstance(node, InlineCode):
            run = paragraph.add_run(node.value)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif isinstance(node, Break):
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif isinstance(node, Link):
            add_hyperlink(paragraph, node.url, plain_text(node.children) or node.url)
        elif isinstance(node, Image):
            paragraph.add_run(f'<image: {image_label(node.url)}>')
        elif isinstance(node, FootnoteReference):
            paragraph.add_run(f'[{node.label or node.identifier}]')
        elif isinstance(node, Html):
            paragraph.add_run(node.value)
        elif isinstance(node, Parent):
            self.render_inlines(paragraph, cast(list[Node], node.children))

    def render_styled_children(
        self, node: Parent, paragraph: DocxParagraph, bold: bool = False, italic: bool = False, strike: bool = False
    ) -> None:
        for child in node.children:
            self.render_inline(child, paragraph, bold=bold, italic=italic, strike=strike)


def list_style_name(base: str, level: int) -> str:
    if level <= 0:
        return base
    return f'{base} {min(level + 1, 5)}'


def add_hyperlink(paragraph: DocxParagraph, href: str, text: str) -> None:
    rel_id = paragraph.part.relate_to(href, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rel_id)

    subrun = paragraph.add_run()
    properties = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000EE')
    properties.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    properties.append(underline)

    subrun._r.append(properties)
    subrun._r.text = text
    hyperlink.append(subrun._r)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph: DocxParagraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    properties.insert_element_before(
        borders,
        'w:shd',
        'w:tabs',
        'w:suppressAutoHyphens',
        'w:kinsoku',
        'w:wordWrap',
        'w:overflowPunct',
        'w:topLinePunct',
        'w:autoSpaceDE',
        'w:autoSpaceDN',
        'w:bidi',
        'w:adjustRightInd',
        'w:snapToGrid',
        'w:spacing',
        'w:ind',
        'w:contextualSpacing',
        'w:mirrorIndents',
        'w:suppressOverlap',
        'w:jc',
        'w:textDirection',
        'w:textAlignment',
        'w:textboxTightWrap',
        'w:outlineLvl',
        'w:divId',
        'w:cnfStyle',
        'w:rPr',
        'w:sectPr',
        'w:pPrChange',
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    borders.append(bottom)


def image_label(url: str) -> str:
    parsed = urlparse(url)
    if parsed.scheme and parsed.netloc and parsed.path:
        return url
    filename = Path(parsed.path or url).name
    return filename or url


def plain_text(children: list[Node]) -> str:
    parts: list[str] = []
    for child in children:
        if isinstance(child, Text | InlineCode):
            parts.append(child.value)
        elif isinstance(child, Image):
            parts.append(child.alt)
        elif isinstance(child, Break):
            parts.append('\n')
        else:
            nested = getattr(child, 'children', None)
            if isinstance(nested, list):
                parts.append(plain_text(cast(list[Node], nested)))
    return ''.join(parts)


@DOCXRenderer.register('root')
def render_root(renderer: DOCXRenderer, node: Root, context: DOCXRenderContext) -> str:
    renderer.render_children(node.children, context)
    return ''


@DOCXRenderer.register('heading')
def render_heading(renderer: DOCXRenderer, node: Heading, context: DOCXRenderContext) -> str:
    paragraph = context.document.add_heading(level=min(node.depth, 9))
    renderer.render_inlines(paragraph, node.children)
    return ''


@DOCXRenderer.register('paragraph')
def render_paragraph(renderer: DOCXRenderer, node: Paragraph, context: DOCXRenderContext) -> str:
    renderer.add_paragraph(context, node.children)
    return ''


@DOCXRenderer.register('list')
def render_list(renderer: DOCXRenderer, node: List, context: DOCXRenderContext) -> str:
    style = list_style_name('List Number' if node.ordered else 'List Bullet', context.list_level)
    for child in node.children:
        if isinstance(child, ListItem):
            renderer.render_list_item(child, style, context)
    return ''


@DOCXRenderer.register('blockquote')
def render_blockquote(renderer: DOCXRenderer, node: Blockquote, context: DOCXRenderContext) -> str:
    for child in node.children:
        if isinstance(child, Paragraph):
            renderer.add_paragraph(context, child.children, style='Quote')
        else:
            renderer.render_node(child, context)
    return ''


@DOCXRenderer.register('code')
def render_code(renderer: DOCXRenderer, node: Code, context: DOCXRenderContext) -> str:
    paragraph = context.document.add_paragraph()
    run = paragraph.add_run(node.value)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return ''


@DOCXRenderer.register('table')
def render_table(renderer: DOCXRenderer, node: Table, context: DOCXRenderContext) -> str:
    rows = [row for row in node.children if isinstance(row, TableRow)]
    if not rows:
        return ''

    column_count = max((len(row.children) for row in rows), default=0)
    table = context.document.add_table(rows=len(rows), cols=column_count)
    if renderer.table_style:
        try:
            table.style = renderer.table_style
        except KeyError as error:
            raise ValueError(f'Unable to apply style {renderer.table_style}.') from error

    for row_index, row_node in enumerate(rows):
        for column_index, cell_node in enumerate(row_node.children):
            if not isinstance(cell_node, TableCell):
                continue
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            renderer.render_cell_content(cell_node, cell, paragraph)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    return ''


@DOCXRenderer.register('thematicBreak')
def render_thematic_break(renderer: DOCXRenderer, node: ThematicBreak, context: DOCXRenderContext) -> str:
    paragraph = context.document.add_paragraph()
    add_bottom_border(paragraph)
    return ''


@DOCXRenderer.register('html')
def render_html(renderer: DOCXRenderer, node: Html, context: DOCXRenderContext) -> str:
    context.document.add_paragraph(node.value)
    return ''


@DOCXRenderer.register('footnoteDefinition')
def render_footnote_definition(renderer: DOCXRenderer, node: FootnoteDefinition, context: DOCXRenderContext) -> str:
    paragraph = context.document.add_paragraph()
    paragraph.add_run(f'[{node.label or node.identifier}] ').bold = True
    if node.children:
        first, *remaining = node.children
        if isinstance(first, Paragraph):
            renderer.render_inlines(paragraph, first.children)
        else:
            renderer.render_node(first, context)
        renderer.render_children(remaining, context)
    return ''
